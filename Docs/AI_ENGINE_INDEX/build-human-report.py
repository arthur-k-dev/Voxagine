"""Build the human-readable Voxagine engine review DOCX.

Design contract:
- documents skill preset: compact_reference_guide
- first-page pattern: editorial_cover
- named overrides: 30 pt editorial title, lead risk callout, compact finding metadata
- canonical content: index.json, review-metadata.json, findings.jsonl
"""

from __future__ import annotations

import json
import math
import tempfile
from collections import Counter, defaultdict
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
INDEX_PATH = HERE / "index.json"
METADATA_PATH = HERE / "review-metadata.json"
FINDINGS_PATH = HERE / "findings.jsonl"
OUTPUT_PATH = HERE / "Voxagine_Engine_Review.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

COLORS = {
    "navy": "0B2545",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "muted": "596675",
    "light_fill": "F4F6F9",
    "table_fill": "E8EEF5",
    "table_alt": "F7F9FB",
    "border": "D5DEE8",
    "risk": "9B1C1C",
    "p1": "C66A12",
    "p2": "2E74B5",
    "p3": "7A8793",
    "gold": "7A5A00",
    "green": "236B4E",
    "white": "FFFFFF",
    "black": "111111",
}

PRIORITY_ORDER = {"P0": 0, "P1": 1, "P2": 2, "P3": 3}
SUBSYSTEM_ORDER = [
    "lifecycle",
    "world",
    "ecs",
    "events",
    "job-system",
    "logging",
    "serialization",
    "resources",
    "physics",
    "rendering",
    "chunks",
    "pathfinding",
    "input",
    "audio",
    "memory",
    "filesystem",
    "player-prefs",
    "time",
    "platform",
    "editor",
    "configuration",
    "build",
    "tests",
]

SUBSYSTEM_LABELS = {
    "lifecycle": "Application lifecycle",
    "world": "World management",
    "ecs": "Entity-component system",
    "events": "Events",
    "job-system": "Jobs and worker coordination",
    "logging": "Logging",
    "serialization": "Serialization and reflection",
    "resources": "Resources and VOX assets",
    "physics": "Physics and voxel grid",
    "rendering": "Rendering and voxel ownership",
    "chunks": "Chunk streaming and far field",
    "pathfinding": "Pathfinding and crowd simulation",
    "input": "Input",
    "audio": "Audio",
    "memory": "Memory allocators",
    "filesystem": "Filesystem and asynchronous I/O",
    "player-prefs": "Player preferences",
    "time": "Time and date",
    "platform": "Platform portability",
    "editor": "Editor",
    "configuration": "Configuration",
    "build": "Build system",
    "tests": "Test and analysis coverage",
}


def load_data():
    index = json.loads(INDEX_PATH.read_text(encoding="utf-8-sig"))
    metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8-sig"))
    findings = []
    for raw_line in FINDINGS_PATH.read_text(encoding="utf-8-sig").splitlines():
        if raw_line.strip():
            findings.append(json.loads(raw_line))
    if len(findings) != index["findings"]["count"]:
        raise RuntimeError("index.json finding count does not match findings.jsonl")
    return index, metadata, findings


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    fonts = run._element.rPr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, color="D5DEE8", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA} DXA: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440.0)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_borders(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_paragraph_shading_and_border(paragraph, fill: str, border_color: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    borders.append(left)
    p_pr.append(borders)


def add_page_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    set_run_font(run, size=8.5, color=COLORS["muted"])


def add_numbering_definition(document: Document, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
        if element.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
        if element.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(2):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), f"%{level + 1}." if ordered else ("•" if level == 0 else "–"))
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)

        text_indent = 540 + level * 360
        hanging = 270
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(text_indent))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(text_indent))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)

        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int = 0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])


def configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(COLORS["black"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(COLORS["navy"])
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing = 1.0

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = rgb(COLORS["dark_blue"])
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.line_spacing = 1.1

    heading_specs = {
        "Heading 1": (16, COLORS["blue"], 18, 10),
        "Heading 2": (13, COLORS["blue"], 14, 7),
        "Heading 3": (12, COLORS["dark_blue"], 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True

    if "Kicker" not in document.styles:
        kicker = document.styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        kicker = document.styles["Kicker"]
    kicker.font.name = "Calibri"
    kicker._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    kicker._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    kicker.font.size = Pt(10)
    kicker.font.bold = True
    kicker.font.color.rgb = rgb(COLORS["gold"])
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(16)
    kicker.paragraph_format.keep_with_next = True

    if "Finding Metadata" not in document.styles:
        meta = document.styles.add_style("Finding Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = document.styles["Finding Metadata"]
    meta.font.name = "Calibri"
    meta._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    meta._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    meta.font.size = Pt(9)
    meta.font.color.rgb = rgb(COLORS["muted"])
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(5)
    meta.paragraph_format.line_spacing = 1.0
    meta.paragraph_format.keep_with_next = True

    document.core_properties.title = "Voxagine Engine Review"
    document.core_properties.subject = "Architecture, correctness, risk, and verification review"
    document.core_properties.author = "Codex"
    document.core_properties.keywords = "Voxagine, engine review, voxel engine, static analysis, architecture"
    document.core_properties.comments = "Generated from Docs/AI_ENGINE_INDEX canonical data."

    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    return section


def configure_header_footer(section, short_revision: str):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = paragraph.add_run("VOXAGINE ENGINE REVIEW")
    set_run_font(left, size=8.5, color=COLORS["muted"], bold=True)
    right = paragraph.add_run(f"\tTECHNICAL AUDIT / {short_revision}")
    set_run_font(right, size=8.5, color=COLORS["muted"])

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Voxagine Engine Review  |  Page ")
    set_run_font(run, size=8.5, color=COLORS["muted"])
    add_page_field(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    set_run_font(run, size=8.5, color=COLORS["muted"])
    add_page_field(paragraph, "NUMPAGES")

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_paragraph = first_footer.paragraphs[0]
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = first_paragraph.add_run(f"Reviewed baseline {short_revision}  |  2026-08-09")
    set_run_font(run, size=8.5, color=COLORS["muted"])


def add_body(document, text: str, bold_lead: str | None = None, keep=False):
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    paragraph.paragraph_format.keep_together = keep
    return paragraph


def add_bullet(document, text: str, bullet_num_id: int, level=0, bold_lead: str | None = None):
    paragraph = document.add_paragraph()
    apply_numbering(paragraph, bullet_num_id, level)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_numbered(document, text: str, number_num_id: int, level=0, bold_lead: str | None = None):
    paragraph = document.add_paragraph()
    apply_numbering(paragraph, number_num_id, level)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_code_path_paragraph(document, label: str, path: str, explanation: str):
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True)
    path_run = paragraph.add_run(path)
    set_run_font(path_run, name="Consolas", size=9, color=COLORS["dark_blue"])
    explanation_run = paragraph.add_run(f" — {explanation}")
    set_run_font(explanation_run)
    return paragraph


def add_lead_callout(document, label: str, text: str, color: str = "9B1C1C"):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.keep_together = True
    set_paragraph_shading_and_border(paragraph, COLORS["light_fill"], color)
    lead = paragraph.add_run(f"{label}: ")
    set_run_font(lead, size=11, color=color, bold=True)
    body = paragraph.add_run(text)
    set_run_font(body, size=11, color=COLORS["black"])
    return paragraph


def add_metric_table(document, metrics):
    table = document.add_table(rows=2, cols=len(metrics))
    widths = [CONTENT_WIDTH_DXA // len(metrics)] * len(metrics)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths)
    for index, (value, label, color) in enumerate(metrics):
        header_cell = table.cell(0, index)
        set_cell_shading(header_cell, COLORS["table_fill"])
        header_paragraph = header_cell.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_paragraph.paragraph_format.space_before = Pt(2)
        header_paragraph.paragraph_format.space_after = Pt(2)
        label_run = header_paragraph.add_run(label)
        set_run_font(label_run, size=8.5, color=COLORS["muted"], bold=True)

        value_cell = table.cell(1, index)
        set_cell_shading(value_cell, "F7F9FB")
        value_paragraph = value_cell.paragraphs[0]
        value_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_paragraph.paragraph_format.space_before = Pt(3)
        value_paragraph.paragraph_format.space_after = Pt(3)
        value_run = value_paragraph.add_run(str(value))
        set_run_font(value_run, size=18, color=color, bold=True)
    repeat_table_header(table.rows[0])
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_verification_table(document, verification):
    table = document.add_table(rows=1, cols=3)
    table.style = None
    headers = ["Configuration", "Result", "Evidence / notes"]
    widths = [2160, 1800, 5400]
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, COLORS["table_fill"])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index < 2 else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        set_run_font(run, size=9.5, color=COLORS["navy"], bold=True)
    for item in verification:
        cells = table.add_row().cells
        values = [item["name"], item["status"], item.get("notes", "")]
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index < 2 else WD_ALIGN_PARAGRAPH.LEFT
            color = COLORS["green"] if str(item["status"]).startswith("passed") else COLORS["muted"]
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9, color=color if index == 1 else COLORS["black"], bold=(index == 1))
    set_table_geometry(table, widths)
    repeat_table_header(table.rows[0])
    document.add_paragraph()
    return table


def find_font(size: int, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def build_risk_chart(findings, output_path: Path):
    grouped = defaultdict(Counter)
    for finding in findings:
        grouped[finding["subsystem"]][finding["priority"]] += 1
    ordered = sorted(grouped.items(), key=lambda item: (-sum(item[1].values()), item[0]))[:14]
    maximum = max(sum(counter.values()) for _, counter in ordered)

    width, height = 1600, 960
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = find_font(42, bold=True)
    label_font = find_font(26, bold=False)
    small_font = find_font(22, bold=True)
    legend_font = find_font(22, bold=False)

    draw.text((60, 42), "Open findings by subsystem and remediation priority", font=title_font, fill="#0B2545")
    draw.text((60, 98), "Counts are stable finding records, not raw defect occurrences.", font=legend_font, fill="#596675")

    legend_x = 760
    for priority in ("P0", "P1", "P2", "P3"):
        color = f"#{COLORS[priority.lower() if priority != 'P0' else 'risk']}" if priority != "P1" else f"#{COLORS['p1']}"
        if priority == "P2":
            color = f"#{COLORS['p2']}"
        if priority == "P3":
            color = f"#{COLORS['p3']}"
        draw.rounded_rectangle((legend_x, 102, legend_x + 30, 132), radius=4, fill=color)
        draw.text((legend_x + 40, 101), priority, font=legend_font, fill="#111111")
        legend_x += 120

    left = 330
    chart_right = 1510
    top = 170
    row_height = 52
    bar_height = 30
    scale = (chart_right - left) / maximum
    color_map = {
        "P0": f"#{COLORS['risk']}",
        "P1": f"#{COLORS['p1']}",
        "P2": f"#{COLORS['p2']}",
        "P3": f"#{COLORS['p3']}",
    }

    for row, (subsystem, counter) in enumerate(ordered):
        y = top + row * row_height
        label = SUBSYSTEM_LABELS.get(subsystem, subsystem.replace("-", " ").title())
        draw.text((55, y - 1), label, font=label_font, fill="#111111")
        x = left
        total = sum(counter.values())
        for priority in ("P0", "P1", "P2", "P3"):
            count = counter.get(priority, 0)
            if count == 0:
                continue
            segment_width = count * scale
            draw.rectangle((x, y, x + segment_width, y + bar_height), fill=color_map[priority])
            if segment_width > 35:
                text = str(count)
                bbox = draw.textbbox((0, 0), text, font=small_font)
                tx = x + (segment_width - (bbox[2] - bbox[0])) / 2
                draw.text((tx, y + 2), text, font=small_font, fill="white")
            x += segment_width
        draw.text((x + 12, y + 1), str(total), font=small_font, fill="#596675")

    image.save(output_path, format="PNG", optimize=True)


def add_picture_with_alt(document, image_path: Path, alt_text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(6.5))
    inline_shape._inline.docPr.set("descr", alt_text)
    return inline_shape


def start_major_section(document, title: str, page_break=True):
    if page_break:
        document.add_page_break()
    heading = document.add_heading(title, level=1)
    heading.paragraph_format.keep_with_next = True
    return heading


def add_cover(document, index):
    revision = index["repository"]["review_baseline_revision"]
    short_revision = revision[:7]

    kicker = document.add_paragraph(style="Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(94)
    kicker.add_run("TECHNICAL AUDIT / ENGINE REVIEW")

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Voxagine Engine Review")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Architecture, correctness, concurrency, resilience, and verification")

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(12)
    metadata.paragraph_format.space_after = Pt(30)
    metadata.paragraph_format.line_spacing = 1.25
    first = metadata.add_run(f"Reviewed revision {short_revision}\n")
    set_run_font(first, size=11, color=COLORS["navy"], bold=True)
    second = metadata.add_run("Full static review completed 2026-08-09\n")
    set_run_font(second, size=10, color=COLORS["muted"])
    third = metadata.add_run("Game Release and full Editor Release builds verified")
    set_run_font(third, size=10, color=COLORS["green"], bold=True)

    add_lead_callout(
        document,
        "Bottom line",
        "The packed voxel/owner pass is a strong foundation. The engine's highest remaining exposure is lifetime coordination across raw-owned objects and asynchronous jobs, followed by malformed-input boundaries and several deterministic editor, input, and allocator faults.",
    )

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    run = note.add_run("Prepared from the versioned AI index in Docs/AI_ENGINE_INDEX")
    set_run_font(run, size=9.5, color=COLORS["muted"], italic=True)


def add_executive_summary(document, index, findings, bullet_num_id):
    start_major_section(document, "Executive summary", page_break=True)

    add_lead_callout(
        document,
        "Overall risk posture: high",
        "Twenty P0 records cover use-after-free, memory corruption, unbounded work, process termination, and parser-driven out-of-bounds paths. Most are concentrated at asynchronous ownership boundaries rather than in the new four-byte voxel representation itself.",
    )

    counts = index["findings"]
    metrics = [
        (counts["count"], "STABLE FINDINGS", COLORS["navy"]),
        (counts["by_priority"]["P0"], "P0 REMEDIATION", COLORS["risk"]),
        (counts["by_status"]["concurrency-risk"], "CONCURRENCY RISKS", COLORS["p1"]),
        (index["sources"]["count"], "INDEXED FILES", COLORS["blue"]),
    ]
    add_metric_table(document, metrics)

    add_body(
        document,
        "Voxagine has a coherent high-level architecture: Application composes process services; Platform selects SDL/Vulkan/audio implementations; WorldManager owns deferred world transitions; World stages ECS mutation; and RTTR connects serialization with editor property handling. The pulled revision substantially improves voxel memory density and adds owner-sidecar, brick occupancy, far-field, range, NaN, and codec verification safeguards.",
    )
    add_body(
        document,
        "The largest systemic weakness is that lifetime is implicit. Jobs, events, resources, editor commands, and world systems frequently retain raw pointers without an owning token or joinable cancellation boundary. Shutdown currently stops jobs and logging before destroying the editor/world consumers that still rely on them. Parser code then forms the second risk cluster: JSON, VOX, and chunk data can reach allocations or indexed access without a complete schema/count/bounds transaction.",
    )

    document.add_heading("Priority conclusions", level=2)
    conclusions = [
        "Fix job-system worker-count and teardown ordering before expanding asynchronous work.",
        "Make world, chunk, pathfinding, physics-integrity, filesystem, and editor callbacks generation/lifetime aware.",
        "Validate JSON, VOX, RLE, settings, dimensions, and arithmetic products before allocating or mutating live state.",
        "Repair deterministic supported-path defects in input controller routing, dynamic ImGui text, command history, playlists, world stack transitions, and voxel accessors.",
        "Quarantine custom allocators until their invariants are replaced and sanitizer/property tests are active.",
        "Promote the successful full game/editor builds into CI, activate the legacy test sources, and add sanitizer/fuzz/stress gates.",
    ]
    for conclusion in conclusions:
        add_bullet(document, conclusion, bullet_num_id)

    document.add_heading("Highest-risk records", level=2)
    top_ids = [
        "VX-LIFE-001",
        "VX-JOB-001",
        "VX-JOB-002",
        "VX-SER-001",
        "VX-SER-002",
        "VX-VOX-002",
        "VX-VOX-004",
        "VX-PHY-002",
        "VX-CHUNK-001",
        "VX-CHUNK-002",
        "VX-PATH-001",
        "VX-PATH-002",
        "VX-INPUT-001",
        "VX-MEM-001",
        "VX-EDITOR-001",
    ]
    finding_map = {item["id"]: item for item in findings}
    for finding_id in top_ids:
        finding = finding_map[finding_id]
        add_bullet(
            document,
            f"{finding_id}: {finding['title']} — {finding['impact']}",
            bullet_num_id,
            bold_lead=f"{finding_id}: ",
        )

    document.add_heading("Strengths to preserve", level=2)
    strengths = [
        "Deferred world swaps are applied at a deliberate GPU synchronization point.",
        "Entity and component mutation is staged rather than modifying active vectors during ordinary ticks.",
        "Platform backends isolate SDL, Vulkan, FMOD, null-audio, and ImGui integration.",
        "RTTR supplies a shared serialization and editor type model.",
        "The pulled CPU voxel is four bytes, with identity moved to a compact uint16 owner sidecar.",
        "Brick occupancy and far-field volume make sparse and distant work explicit.",
        "Recent voxel/chunk code includes range, NaN, owner, and RLE round-trip diagnostics.",
        "Normal resource unload precedes platform teardown, preserving backend availability for common destruction.",
    ]
    for strength in strengths:
        add_bullet(document, strength, bullet_num_id)

    document.add_heading("Verification outcome", level=2)
    add_verification_table(document, index["verification"])
    add_body(
        document,
        "The game Release build produced voxagine_vulkan.lib, voxagine.lib, rttr_core.dll, and BitBuster.exe. A separate fresh VOXAGINE_BUILD_EDITOR=ON Release configuration compiled and linked the full editor engine plus BitBuster. The editor build is successful but not warning-clean; warning debt is recorded as VX-BUILD-005, and the concrete missing return in VoxFrame::GetVoxelColor is VX-VOX-006.",
    )


def add_architecture(document, bullet_num_id):
    start_major_section(document, "Architecture and engineering model")
    add_body(
        document,
        "Application::Run is the process composition root. It initializes filesystem, logging, serialization, settings, jobs, platform services, game state, and the optional editor. The frame loop applies deferred world swaps, polls input/UI, processes game and job completions, ticks the active world, advances fixed simulation, uploads camera data, renders, and presents.",
    )
    add_body(
        document,
        "WorldManager owns a stack of raw World pointers. Mutations are queued and applied at the start of a later frame after a GPU wait. Each World owns staged entities, ordinary component systems, and a separate RenderSystem. Script, physics, audio, chunk, and rendering work share the world's JobQueue route; pathfinding entities schedule additional jobs through the same application service.",
    )

    document.add_heading("Runtime layers", level=2)
    layers = [
        "Application layer — lifecycle, process services, game hooks, and frame sequencing.",
        "Platform layer — SDL window/input, Vulkan render context, ImGui, timers, and FMOD or NullAudioContext.",
        "World/ECS layer — deferred world stack, staged entities/components, systems, camera, hierarchy, and reflection-visible state.",
        "Voxel layer — VoxModel frames, VoxelBaker, mapped voxel/color/owner buffers, brick occupancy, physics VoxelGrid, chunk RLE, and far-field volume.",
        "Tooling layer — editor world/camera, hierarchy, inspector, property renderers, asset views, logging console, and undo/redo.",
    ]
    for layer in layers:
        lead = layer.split(" — ", 1)[0] + " — "
        add_bullet(document, layer, bullet_num_id, bold_lead=lead)

    document.add_heading("Frame and completion boundary", level=2)
    sequence = [
        "Wait for the GPU and apply deferred world operations when required.",
        "Poll the window, input, and ImGui, then clear the render context.",
        "Run game OnUpdate and process main-thread Job::Finish callbacks.",
        "Run world pre-tick/tick, bounded fixed steps, and post-tick.",
        "Upload optional camera state; render fixed data, gizmos, editor, and game draw hooks.",
        "Present through the Vulkan render context.",
    ]
    for item in sequence:
        add_bullet(document, item, bullet_num_id)

    document.add_heading("Ownership invariant that needs to become explicit", level=2)
    add_lead_callout(
        document,
        "Required invariant",
        "No asynchronous job, event callback, resource-release callback, or undo command may outlive the object graph it references. Today that rule is convention-based; the highest-priority roadmap makes it enforceable with owned results, connection tokens, generations, cancellation, and join barriers.",
        color=COLORS["p1"],
    )

    document.add_heading("Voxel pipeline at the reviewed revision", level=2)
    add_body(
        document,
        "VoxRenderer selects a VoxModel frame and transform. VoxelBaker expands compact solid positions/colors into the global mapped grid, associates a compact owner slot, and marks occupancy. Physics and chunk systems read or modify related voxel/chunk structures, while the far-field path maintains a quarter-resolution representation for distant geometry. The owner sidecar supports 65,533 distinct static owners and deliberately does not recycle slots, preventing stale identity reuse at the cost of a finite long-session budget.",
    )


def add_risk_profile(document, findings):
    start_major_section(document, "Risk profile")
    add_body(
        document,
        "The chart shows stable finding records by subsystem and priority. It should be used for sequencing, not as a defect-density metric: one record can group several tightly related observations, and severity is determined by consequence rather than file count.",
    )
    with tempfile.TemporaryDirectory(prefix="voxagine-report-") as temp_dir:
        chart_path = Path(temp_dir) / "risk-profile.png"
        build_risk_chart(findings, chart_path)
        add_picture_with_alt(
            document,
            chart_path,
            "Stacked horizontal bar chart of Voxagine review findings by subsystem and P0 through P3 priority.",
        )
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    run = caption.add_run("Figure 1. Stable findings by subsystem and remediation priority")
    set_run_font(run, size=9, color=COLORS["muted"], italic=True)

    document.add_heading("Evidence interpretation", level=2)
    evidence_notes = [
        "Confirmed — deterministic from the reviewed source or compiler output.",
        "Concurrency-risk — an unsynchronized interleaving exists; stress/sanitizer evidence is still desirable.",
        "Needs-runtime-validation — source-supported but dependent on runtime/API behavior.",
        "Coverage-gap — missing verification, not a direct code defect.",
        "Limitation — an intentional or architectural constraint future work must preserve or consciously redesign.",
    ]
    for note in evidence_notes:
        lead = note.split(" — ", 1)[0] + " — "
        add_bullet(document, note, document._bullet_num_id, bold_lead=lead)


def add_roadmap(document, bullet_num_id, number_num_id):
    start_major_section(document, "Recommended remediation roadmap")
    add_body(
        document,
        "The roadmap follows dependency order. Closing parser symptoms before lifetime coordination, or adding more asynchronous work before cancellation is joinable, would leave the most consequential failure modes intact.",
    )

    phases = [
        (
            "Phase 1 — Establish safe lifetime boundaries",
            [
                "Clamp worker creation and redesign JobManager deinitialization so workers join before queues are deleted.",
                "Define exactly-once job terminal behavior for success, failure, pending cancellation, and running cancellation.",
                "Reorder Application shutdown around stopped producers, joined world work, editor/world teardown, resource unload, then global services.",
                "Add generation/lifetime tokens for world, physics integrity, chunks, pathfinding, filesystem, editor events, and completion callbacks.",
            ],
        ),
        (
            "Phase 2 — Harden data and arithmetic boundaries",
            [
                "Add transactional schemas for generic settings, worlds, prefabs, far-field model paths, and RTTR component construction.",
                "Make VOX and RLE parsers count-, chunk-, size-, and allocation-safe; fuzz them under ASan/UBSan.",
                "Validate settings once before subsystem construction, including positive timers, nonzero dimensions, divisibility, and checked products.",
                "Fix solid-count, empty-frame, GetVoxelColor, spherical-radius, grid-boundary, and stamp-transform behavior.",
            ],
        ),
        (
            "Phase 3 — Repair deterministic runtime/editor paths",
            [
                "Fix specific-player controller indexing and finish action/axis/map lifecycle symmetry.",
                "Replace all dynamic ImGui format strings with unformatted or constant-format calls.",
                "Move editor command history to owned commands and stable target resolution; disconnect events during teardown.",
                "Fix playlist, AudioSource/FMOD, main-camera optionality, world stack guards, timer catch-up, and POSIX mode/write behavior.",
                "Remove or quarantine the custom allocator family until redesigned and proven.",
            ],
        ),
        (
            "Phase 4 — Make the repaired behavior continuously verifiable",
            [
                "Compile full game and editor Debug/Release configurations in CI on supported platforms.",
                "Integrate legacy GoogleTests with CMake/CTest and add malformed asset, teardown, chunk churn, and path stress suites.",
                "Add ASan/UBSan, TSAN where supported, fuzzers, and static analysis.",
                "Eliminate the first-party warning baseline and stage warnings-as-errors for C++ and shaders.",
                "Regenerate the AI index and resolve stable IDs as fixes land; never delete history without recording resolution evidence.",
            ],
        ),
    ]
    for title, actions in phases:
        document.add_heading(title, level=2)
        for action in actions:
            add_bullet(document, action, bullet_num_id)

    document.add_heading("Definition of done for a finding", level=2)
    steps = [
        "The current code no longer contains the evidenced path, including all grouped variants.",
        "A focused regression test fails on the reviewed baseline and passes with the fix.",
        "Relevant game/editor builds and sanitizer/stress/fuzz checks pass.",
        "The existing stable finding ID is updated with resolution revision and validation evidence.",
        "The generated index reports matching baseline/source hashes and valid JSONL.",
    ]
    for step in steps:
        add_numbered(document, step, number_num_id)


def add_scope_and_index(document, index, metadata, bullet_num_id):
    start_major_section(document, "Scope, method, and durable index")
    add_body(
        document,
        "This was a deep static review plus compile verification. It traced architecture, lifecycle, ownership, jobs, ECS, serialization, resources, rendering, voxel data, physics, chunks, pathfinding, input, audio, editor, memory, filesystem, timing, build topology, and test coverage. Vendored implementations and binary assets were excluded except at their integration boundaries.",
    )

    document.add_heading("Included surfaces", level=2)
    for item in metadata["scope"]["included"]:
        add_bullet(document, item, bullet_num_id)

    document.add_heading("Excluded surfaces", level=2)
    for item in metadata["scope"]["excluded"]:
        add_bullet(document, item, bullet_num_id)

    document.add_heading("Review limits", level=2)
    add_body(
        document,
        "The review did not launch a graphical Vulkan/FMOD session, fuzz untrusted assets, run thread sanitizers, or execute every game/editor path. Concurrency records identify unsafe interleavings but still benefit from stress/TSAN reproduction. Existing legacy tests were not runnable because they are not integrated with root CMake. These limitations are findings, not hidden assumptions.",
    )
    add_body(document, metadata["review"]["working_tree_note"])

    document.add_heading("Durable artifact routes", level=2)
    routes = [
        ("Machine entry", "Docs/AI_ENGINE_INDEX/index.json", "baseline, counts, verification, hashes, and artifact routes"),
        ("Finding database", "Docs/AI_ENGINE_INDEX/findings.jsonl", "one stable issue record per line"),
        ("Source manifest", "Docs/AI_ENGINE_INDEX/source-manifest.jsonl", "one first-party file with symbols, includes, lines, bytes, and SHA-256 per line"),
        ("Architecture", "Docs/AI_ENGINE_INDEX/ARCHITECTURE.md", "lifecycle, ownership, threading, and subsystem model"),
        ("Source map", "Docs/AI_ENGINE_INDEX/SOURCE_MAP.md", "task and symbol routing"),
        ("Human findings", "Docs/AI_ENGINE_INDEX/FINDINGS.md", "version-controlled readable explanation"),
        ("Refresh tool", "Docs/AI_ENGINE_INDEX/refresh-index.ps1", "regenerates and validates the central index"),
    ]
    for label, path, explanation in routes:
        add_code_path_paragraph(document, label, path, explanation)

    document.add_heading("Efficient lookup protocol", level=2)
    lookup_steps = [
        "Read index.json first and confirm review_baseline_revision matches the intended source revision.",
        "Filter findings.jsonl by stable ID, subsystem, priority, severity, status, evidence path, or symbol.",
        "Use source-manifest.jsonl when the current implementation location is unknown.",
        "Open only the routed current source before fixing; reviewed line hints can drift.",
        "Regenerate after source/finding changes and preserve stable IDs across resolution.",
    ]
    for step in lookup_steps:
        add_numbered(document, step, document._number_num_id)


def add_finding_entry(document, finding, bullet_num_id):
    heading = document.add_heading(f"{finding['id']} — {finding['title']}", level=3)
    heading.paragraph_format.keep_with_next = True

    metadata = document.add_paragraph(style="Finding Metadata")
    priority_color = COLORS["risk"] if finding["priority"] == "P0" else COLORS.get(finding["priority"].lower(), COLORS["muted"])
    if finding["priority"] == "P1":
        priority_color = COLORS["p1"]
    priority = metadata.add_run(finding["priority"])
    set_run_font(priority, size=9, color=priority_color, bold=True)
    separators = [
        f"  |  severity: {finding['severity']}",
        f"  |  evidence: {finding['status']}",
        f"  |  confidence: {finding['confidence']}",
        f"  |  subsystem: {finding['subsystem']}",
    ]
    for text in separators:
        run = metadata.add_run(text)
        set_run_font(run, size=9, color=COLORS["muted"])

    summary = add_body(document, finding["summary"], keep=True)
    summary.paragraph_format.keep_with_next = True

    evidence_label = document.add_paragraph()
    evidence_label.paragraph_format.space_after = Pt(2)
    evidence_label.paragraph_format.keep_with_next = True
    run = evidence_label.add_run("Evidence")
    set_run_font(run, size=10, color=COLORS["navy"], bold=True)
    for item in finding["evidence"]:
        path = item.get("path", "")
        symbol = item.get("symbol", "")
        lines = item.get("lines", item.get("line", ""))
        location = path
        if symbol:
            location += f" :: {symbol}"
        if lines:
            location += f" (reviewed lines {lines})"
        note = item.get("note", "")
        paragraph = document.add_paragraph()
        apply_numbering(paragraph, bullet_num_id, 0)
        loc_run = paragraph.add_run(location)
        set_run_font(loc_run, name="Consolas", size=8.5, color=COLORS["dark_blue"], bold=True)
        if note:
            note_run = paragraph.add_run(f" — {note}")
            set_run_font(note_run, size=9.5)

    impact = document.add_paragraph()
    impact.paragraph_format.space_before = Pt(2)
    impact.paragraph_format.space_after = Pt(4)
    label = impact.add_run("Impact. ")
    set_run_font(label, color=COLORS["risk"], bold=True)
    body = impact.add_run(finding["impact"])
    set_run_font(body)

    recommendation = document.add_paragraph()
    recommendation.paragraph_format.space_after = Pt(4)
    label = recommendation.add_run("Recommended correction. ")
    set_run_font(label, color=COLORS["green"], bold=True)
    body = recommendation.add_run(finding["recommendation"])
    set_run_font(body)

    validation_label = document.add_paragraph()
    validation_label.paragraph_format.space_after = Pt(2)
    validation_label.paragraph_format.keep_with_next = True
    run = validation_label.add_run("Validation gate")
    set_run_font(run, size=10, color=COLORS["navy"], bold=True)
    for validation in finding.get("validation", []):
        add_bullet(document, validation, bullet_num_id)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_detailed_findings(document, findings, bullet_num_id):
    start_major_section(document, "Detailed findings")
    add_body(
        document,
        "Every record below is sourced from findings.jsonl and retains its stable ID. P0/P1/P2/P3 indicate remediation order; severity indicates consequence. Paths and symbols are the durable evidence keys, while line numbers are navigation hints for revision 743c4c6.",
    )

    grouped = defaultdict(list)
    for finding in findings:
        grouped[finding["subsystem"]].append(finding)

    known = [name for name in SUBSYSTEM_ORDER if name in grouped]
    remaining = sorted(name for name in grouped if name not in known)
    for subsystem in known + remaining:
        label = SUBSYSTEM_LABELS.get(subsystem, subsystem.replace("-", " ").title())
        count = len(grouped[subsystem])
        heading = document.add_heading(f"{label} ({count})", level=2)
        heading.paragraph_format.keep_with_next = True
        ordered = sorted(grouped[subsystem], key=lambda item: (PRIORITY_ORDER[item["priority"]], item["id"]))
        for finding in ordered:
            add_finding_entry(document, finding, bullet_num_id)


def add_closing(document, index):
    start_major_section(document, "Conclusion")
    add_lead_callout(
        document,
        "Recommended decision",
        "Keep the pulled packed-voxel architecture, but treat asynchronous ownership and untrusted-input validation as the next engine milestone. Do not scale pathfinding, chunk streaming, or background voxel work until cancellation/join and snapshot semantics are explicit and tested.",
        color=COLORS["green"],
    )
    add_body(
        document,
        "The review artifacts are designed to stay useful after this report. Future agents should begin at index.json, filter stable finding records, use source-manifest.jsonl to route to current symbols, and update rather than replace IDs as issues are fixed. That preserves engineering memory while keeping lookup cost low.",
    )
    final = document.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.paragraph_format.space_before = Pt(30)
    run = final.add_run(
        f"End of report  |  {index['findings']['count']} findings  |  {index['sources']['count']} indexed files  |  baseline {index['repository']['review_baseline_revision'][:7]}"
    )
    set_run_font(run, size=9.5, color=COLORS["muted"], italic=True)


def build_report():
    index, metadata, findings = load_data()
    document = Document()
    section = configure_document(document)
    short_revision = index["repository"]["review_baseline_revision"][:7]
    configure_header_footer(section, short_revision)

    bullet_num_id = add_numbering_definition(document, ordered=False)
    number_num_id = add_numbering_definition(document, ordered=True)
    document._bullet_num_id = bullet_num_id
    document._number_num_id = number_num_id

    add_cover(document, index)
    add_executive_summary(document, index, findings, bullet_num_id)
    add_architecture(document, bullet_num_id)
    add_risk_profile(document, findings)
    add_roadmap(document, bullet_num_id, number_num_id)
    add_scope_and_index(document, index, metadata, bullet_num_id)
    add_detailed_findings(document, findings, bullet_num_id)
    add_closing(document, index)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")
    print(f"Findings included: {len(findings)}")


if __name__ == "__main__":
    build_report()
